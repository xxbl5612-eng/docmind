"""Generate AGI Full-Stack Development course project report for DocMind.

Matches the template: XXX项目小组-AGI全栈开发大作业报告（模板本）
Personal information fields left blank for user to fill.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT = "E:/projects/docmind/DocMind_AGI全栈开发大作业报告.docx"
SCREENSHOT_DIR = "E:/projects/docmind/scripts/pptx"

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── Helper functions ──

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, indent=False, size=12, align=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
    # Data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()
    return table


# ═══════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════

for _ in range(3):
    doc.add_paragraph()

add_para("XX学院综合与创新实验报告", bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

cover_items = [
    ("课程名称：", "AGI全栈开发"),
    ("项目组长", "____________    学号 ____________"),
    ("项目成员", "____________    学号 ____________"),
    ("", "____________    学号 ____________"),
    ("题    目", "AGI全栈开发——DocMind智能文档处理平台架构设计及全栈实现（单人项目）"),
    ("指导教师", "刘雨飏"),
    ("学年学期", "2025—2026学年 第二学期"),
    ("实验时间", "____年____月____日——____年____月____日"),
]
for label, value in cover_items:
    p = doc.add_paragraph()
    if label:
        run = p.add_run(label)
        run.bold = True
    p.add_run(value)

doc.add_page_break()

# ═══════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════

add_heading_styled("目  录", level=1)
toc_items = [
    "1  项目计划",
    "    1.1  项目简介",
    "    1.2  项目背景",
    "    1.3  需求分析设计",
    "        1.3.1  痛点分析/核心业务",
    "        1.3.2  产品结构图",
    "        1.3.3  主要业务流程分析",
    "        1.3.4  信息结构图",
    "2  前端设计与开发",
    "    2.1  模块架构设计",
    "        2.1.1  模块目录结构",
    "        2.1.2  页面目录结构",
    "    2.2  核心功能实现",
    "        2.2.1  页面：Dashboard工作台页面",
    "        2.2.2  组件：DraggableAiAssistant组件",
    "        2.2.3  服务：ApiService服务层",
    "        2.2.4  管道：i18n国际化管道",
    "3  微服务设计与开发",
    "    3.1  UML时序图设计",
    "    3.2  关系数据库模式设计",
    "    3.3  范式数据的实体导入",
    "4  难点功能实现",
    "5  项目管理总结",
    "    5.1  项目迭代周期分析",
    "    5.2  项目成员贡献分析",
    "    5.3  项目代码量分析",
    "6  项目产品展示",
]
for item in toc_items:
    add_para(item, size=12)

doc.add_page_break()

# ═══════════════════════════════════════════
# SECTION 1: 项目计划
# ═══════════════════════════════════════════

add_heading_styled("1  项目计划", level=1)

add_heading_styled("1.1  项目简介", level=2)
add_para(
    "项目名称：DocMind——全场景智能文档处理平台（Full-Scenario Intelligent Document Processing Assistant）。"
    "本项目是一个面向企业办公、学术研究和内容创作场景的智能文档处理Web应用，"
    "支持多格式文档上传与解析、AI驱动的文档校对/总结/改写/翻译、基于WebSocket的实时协作编辑、"
    "版本控制与历史追溯、GitHub仓库双向集成，"
    "以及Phase 2新增的PDF导出增强（页码/目录/书签）、OCR图片文字识别（PaddleOCR+EasyOCR双引擎、结构化表格输出、语言自动检测、二维码识别）、"
    "高级语义搜索（BM25关键词+FAISS向量混合搜索、跨文档检索、关键词高亮与RAG问答）。",
    indent=True
)
add_para(
    "技术栈：前端React 19 + TypeScript 5 + Vite 8 + Tailwind CSS 4 + TanStack Query 5；"
    "后端Python 3.12 + FastAPI + SQLAlchemy 2（异步）+ Celery 5；"
    "数据库PostgreSQL 16 + Redis 7 + MinIO；"
    "AI引擎DeepSeek API v1 + BCE Embedding（ONNX）+ FAISS向量检索。",
    indent=True
)

add_heading_styled("1.2  项目背景", level=2)
add_para(
    "随着大语言模型（LLM）技术的快速发展，AGI正逐步渗透到各个行业领域。"
    "传统的文档处理方式面临格式多样难以统一、人工审校效率低易出错、团队协作缺乏实时同步、"
    "文档版本管理混乱、格式排版耗时费力、内容提取与分析困难等痛点。",
    indent=True
)
add_para(
    "在学术支持方面，学生撰写论文时面临文献阅读量大、格式调整繁琐、多人协作低效的问题；"
    "在企业场景中，办公人员日常需要处理工作总结、汇报方案、通知公告等多种文档类型；"
    "对于技术团队，技术文档的协作编写和版本管理是刚需。",
    indent=True
)
add_para(
    "市场现有解决方案（如Google Docs、腾讯文档、飞书文档等）提供了基础协作功能，"
    "但在AI深度处理（校对/总结/改写/翻译）、语义级文档检索、自动化格式转换等方面仍有不足。"
    "DocMind正是基于这些痛点，利用AGI技术栈构建的全场景智能文档处理平台。",
    indent=True
)

add_heading_styled("1.3  需求分析设计", level=2)

add_heading_styled("1.3.1  痛点分析/核心业务", level=3)
add_para("传统文档处理面临以下八大痛点：", indent=True)

pain_points = [
    ["1", "格式不支持", "仅能处理单一格式文件", "多格式引擎：原生支持DOCX/MD/PDF/TXT/PPTX/JPG，自动识别与解析"],
    ["2", "效率低下", "人工校对耗时且易出错", "AI智能校对：基于DeepSeek的语法/拼写/逻辑纠错，秒级响应"],
    ["3", "协作困难", "文件通过邮件/U盘反复传输", "实时协作：WebSocket多人同步编辑，操作日志完整追踪"],
    ["4", "版本混乱", "修改后找不到历史版本", "版本控制：自动保存版本，任意回滚 + Diff对比"],
    ["5", "排版费力", "手动调整格式反复操作", "智能排版：AI自动识别标题层级、段落格式，一键美化"],
    ["6", "内容提取难", "无法快速获取核心信息", "智能总结：AI自动提取关键信息，3秒生成摘要"],
    ["7", "工具割裂", "写作、翻译、改写需切换工具", "一体化平台：9种AI工具集成，支持链式处理"],
    ["8", "搜索低效", "关键词匹配无法理解语义", "语义搜索：BM25+向量混合搜索，跨文档RAG问答"],
]
add_table(["序号", "痛点", "传统方式", "DocMind解决方案"], pain_points)

add_para("核心业务功能矩阵：", bold=True)
core_biz = [
    ["文档管理", "上传/下载/删除/文件夹组织", "RESTful API + MinIO存储"],
    ["AI处理", "校对/总结/改写/扩写/精简/翻译/提取/转换/问答", "DeepSeek API + 策略模式Pipeline"],
    ["实时协作", "多人编辑/光标同步/操作日志", "WebSocket + 操作OT日志"],
    ["版本控制", "自动快照/回滚/Diff对比", "DocumentVersion模型 + 内容差异算法"],
    ["语义搜索", "混合搜索/跨文档搜索/RAG问答", "BCE Embedding + FAISS + BM25"],
    ["OCR识别", "图片/PDF文字识别/表格提取/二维码检测", "PaddleOCR + EasyOCR + Pyzbar"],
    ["PDF导出", "格式转换/压缩/水印/加密/合并/页码", "LibreOffice + Ghostscript + PyPDF2"],
    ["AI助手", "全局浮动对话/文档上下文感知", "DeepSeek Chat API + 可拖拽UI"],
]
add_table(["功能模块", "能力", "技术实现"], core_biz)

add_heading_styled("1.3.2  产品结构图", level=3)
add_para(
    "DocMind产品采用前后端分离的三层架构：",
    indent=True
)
add_para(
    "前端层（React 19 + TypeScript）：包含Landing/Login/Register公开页面，"
    "Dashboard工作台、DocumentEditor编辑器、AiToolPanel AI工具面板、"
    "VersionHistory版本历史、Collaboration协作、GitHubImport导入、"
    "Settings设置、AdminDashboard管理等受保护页面，"
    "以及DocxViewer/MdViewer/PdfViewer/TxtViewer/PptxViewer五种文档查看器，"
    "DraggableAiAssistant全局AI助手组件，SemanticSearch语义搜索组件。",
    indent=True
)
add_para(
    "后端层（FastAPI + Celery）：包含auth/users/documents/processing/versions/"
    "collaboration/github/oauth/admin/operations/ai共12个路由模块，"
    "8个核心服务模块（auth/document/ai/collaboration/github/oauth/user/version），"
    "以及Phase 2新增的ocr_service/search_service/pdf_export_service/pptx_render_service。",
    indent=True
)
add_para(
    "数据层（PostgreSQL 16 + Redis 7 + MinIO）：16张数据表覆盖用户/文档/版本/协作/"
    "OAuth/GitHub/配额/AI任务等全部业务域，Redis提供缓存和消息队列，MinIO提供对象存储。",
    indent=True
)

add_heading_styled("1.3.3  主要业务流程分析", level=3)
add_para("以下使用PlantUML描述核心业务流程：", indent=True)
add_para("【文档上传→AI处理流程】", bold=True)
add_para(
    "@startuml\n"
    "用户 -> 前端: 选择文件上传\n"
    "前端 -> FastAPI: POST /api/v1/documents/upload\n"
    "FastAPI -> 解析服务: 提取文本内容\n"
    "解析服务 -> 数据库: 保存文档元数据\n"
    "解析服务 -> MinIO: 保存原始文件\n"
    "FastAPI -> 搜索服务: auto_index_on_upload\n"
    "搜索服务 -> BCE模型: 文本向量化\n"
    "搜索服务 -> FAISS: 构建语义索引\n"
    "FastAPI --> 前端: 返回文档ID\n"
    "用户 -> 前端: 选择AI操作(校对/总结/改写)\n"
    "前端 -> FastAPI: POST /documents/{id}/ai/proofread\n"
    "FastAPI -> AI服务: 加载文档内容+构建Prompt\n"
    "AI服务 -> DeepSeek API: 发送处理请求\n"
    "DeepSeek API --> AI服务: 返回处理结果\n"
    "AI服务 --> 前端: 展示结果\n"
    "@enduml",
    indent=True, size=10
)

add_para("【语义搜索→RAG问答流程】", bold=True)
add_para(
    "@startuml\n"
    "用户 -> 前端: 输入搜索关键词\n"
    "前端 -> FastAPI: POST /documents/{id}/ai/search (mode=hybrid)\n"
    "FastAPI -> BCE模型: 查询向量化\n"
    "FastAPI -> FAISS: 向量相似度检索(top_k*2)\n"
    "FastAPI -> BM25: 关键词评分\n"
    "FastAPI -> 融合排序: 加权融合(vector*0.7 + bm25*0.3)\n"
    "FastAPI -> 高亮: 提取匹配片段\n"
    "FastAPI --> 前端: 返回排序结果+高亮片段\n"
    "用户 -> 前端: 切换QA模式提问\n"
    "前端 -> FastAPI: POST /documents/{id}/ai/search/qa\n"
    "FastAPI -> 搜索: 检索相关片段(top_k)\n"
    "FastAPI -> DeepSeek: 构建RAG Prompt(上下文+问题)\n"
    "DeepSeek --> FastAPI: 返回答案+引用来源\n"
    "FastAPI --> 前端: 展示答案+来源标注\n"
    "@enduml",
    indent=True, size=10
)

add_heading_styled("1.3.4  信息结构图", level=3)
add_para("系统核心数据实体及关系：", indent=True)
add_para(
    "User（用户）1──N Document（文档）—— owner_id FK\n"
    "Document 1──N DocumentVersion（版本）—— document_id FK\n"
    "Document 1──N CollaborationSession（协作会话）—— document_id FK\n"
    "CollaborationSession 1──N Collaborator（协作者）—— session_id FK\n"
    "Document N──M Tag（标签）—— tag_associations中间表\n"
    "Document 1──N AIProcessingJob（AI任务）—— document_id FK\n"
    "User 1──1 UserQuota（配额）—— user_id FK\n"
    "User 1──N OAuthAccount（OAuth账户）—— user_id FK\n"
    "User 1──N RefreshToken（刷新令牌）—— user_id FK\n"
    "User 1──N OperationLog（操作日志）—— user_id FK",
    indent=True, size=10
)

doc.add_page_break()

# ═══════════════════════════════════════════
# SECTION 2: 前端设计与开发
# ═══════════════════════════════════════════

add_heading_styled("2  前端设计与开发", level=1)

add_heading_styled("2.1  模块架构设计", level=2)

add_heading_styled("2.1.1  模块目录结构", level=3)
add_para("前端项目采用基于功能模块的目录组织方式：", indent=True)
add_para(
    "frontend/src/\n"
    "├── pages/                    # 路由级页面（12个页面）\n"
    "│   ├── Landing.tsx           # 着陆页\n"
    "│   ├── Login.tsx             # 登录页\n"
    "│   ├── Register.tsx          # 注册页\n"
    "│   ├── Dashboard.tsx         # 工作台（文档列表+统计卡片）\n"
    "│   ├── DocumentEditor.tsx    # 文档编辑器\n"
    "│   ├── AiToolPanel.tsx       # AI工具面板\n"
    "│   ├── VersionHistory.tsx    # 版本历史\n"
    "│   ├── Collaboration.tsx     # 协作管理\n"
    "│   ├── GitHubImport.tsx      # GitHub导入\n"
    "│   ├── Settings.tsx          # 用户设置\n"
    "│   ├── AdminDashboard.tsx    # 管理后台\n"
    "│   └── AuthCallback.tsx      # OAuth回调\n"
    "├── components/               # 可复用组件（26个）\n"
    "│   ├── layout/               # 布局组件\n"
    "│   │   ├── AppLayout.tsx     # 主布局（含AI助手注入）\n"
    "│   │   ├── Navbar.tsx        # 导航栏\n"
    "│   │   ├── Sidebar.tsx       # 侧边栏\n"
    "│   │   └── Footer.tsx        # 页脚\n"
    "│   ├── ai/                   # AI组件\n"
    "│   │   └── DraggableAiAssistant.tsx  # 可拖拽AI助手（296行）\n"
    "│   ├── search/               # 搜索组件\n"
    "│   │   └── SemanticSearch.tsx  # 语义搜索面板（172行）\n"
    "│   ├── viewer/               # 文档查看器\n"
    "│   │   ├── DocxViewer.tsx\n"
    "│   │   ├── MdViewer.tsx\n"
    "│   │   ├── PdfViewer.tsx\n"
    "│   │   ├── TxtViewer.tsx\n"
    "│   │   └── PptxViewer.tsx    # PPTX幻灯片查看器（198行）\n"
    "│   ├── ui/                   # 基础UI组件\n"
    "│   │   ├── Button.tsx, Input.tsx, Modal.tsx, Toast.tsx\n"
    "│   │   └── Dropdown.tsx, Card.tsx, Badge.tsx, Spinner.tsx\n"
    "│   ├── common/               # 通用组件\n"
    "│   │   ├── ErrorBoundary.tsx\n"
    "│   │   ├── LoadingSpinner.tsx\n"
    "│   │   └── EmptyState.tsx\n"
    "│   └── github/               # GitHub组件\n"
    "│       ├── RepoBrowser.tsx\n"
    "│       ├── FileTree.tsx\n"
    "│       └── PRCard.tsx\n"
    "├── lib/                      # 工具库\n"
    "│   ├── api.ts                # API客户端（Axios + TanStack Query）\n"
    "│   ├── i18n.ts               # 国际化引擎（zh-CN/en-US）\n"
    "│   └── utils.ts              # 通用工具函数\n"
    "├── types/                    # TypeScript类型定义\n"
    "│   └── index.ts              # 80+接口/类型\n"
    "└── locales/                  # 国际化词条\n"
    "    ├── zh.json               # 中文词条\n"
    "    └── en.json               # 英文词条",
    indent=True, size=9
)

add_heading_styled("2.1.2  页面目录结构", level=3)
add_para("React Router 7路由配置（BrowserRouter模式）：", indent=True)
routes = [
    ["/", "Landing", "公开", "着陆页"],
    ["/login", "Login", "公开", "登录"],
    ["/register", "Register", "公开", "注册"],
    ["/dashboard", "Dashboard", "受保护", "工作台"],
    ["/editor/:id", "DocumentEditor", "受保护", "文档编辑器"],
    ["/ai/:id", "AiToolPanel", "受保护", "AI工具面板"],
    ["/versions/:id", "VersionHistory", "受保护", "版本历史"],
    ["/collaboration/:id", "Collaboration", "受保护", "协作管理"],
    ["/github/import", "GitHubImport", "受保护", "GitHub导入"],
    ["/settings", "Settings", "受保护", "用户设置"],
    ["/admin", "AdminDashboard", "受保护", "管理后台"],
    ["/auth/github/callback", "AuthCallback", "公开", "OAuth回调"],
    ["*", "NotFound", "公开", "404兜底"],
]
add_table(["路径", "组件", "权限", "说明"], routes)

add_heading_styled("2.2  核心功能实现", level=2)

# ── 2.2.1 Dashboard ──
add_heading_styled("2.2.1  页面：Dashboard工作台页面", level=3)
add_para("1) HTML页面模板/DOM结构", bold=True)
add_para(
    "Dashboard页面采用卡片网格布局，顶部为统计卡片行（文档数/AI调用数/存储空间），"
    "下方为文档列表区域（支持全部/我的文档/共享给我三个Tab筛选），"
    "每个文档卡片显示标题、格式标签、大小、状态和时间。"
    "页面嵌入DraggableAiAssistant全局AI助手组件（通过AppLayout注入）。",
    indent=True
)
add_para(
    "DOM结构：\n"
    "AppLayout\n"
    "├── Navbar（导航栏 + 用户菜单）\n"
    "├── main\n"
    "│   ├── StatsCards（统计卡片网格 grid-cols-3）\n"
    "│   ├── TabBar（全部/我的文档/共享给我）\n"
    "│   ├── ActionBar（上传按钮 + 从GitHub导入按钮）\n"
    "│   └── DocumentGrid（文档卡片网格 grid-cols-1 md:grid-cols-2 lg:grid-cols-3）\n"
    "└── DraggableAiAssistant（fixed定位浮动助手）",
    indent=True, size=10
)

add_para("2) SCSS样式与视觉效果", bold=True)
add_para(
    "使用Tailwind CSS v4原子化样式系统，核心样式方案：\n"
    "• 统计卡片：bg-white rounded-xl shadow-sm border border-surface-200，hover时shadow-md过渡\n"
    "• 文档卡片：bg-white rounded-lg border，hover:border-primary-300 hover:shadow-md\n"
    "• 布局：max-w-7xl mx-auto px-4 sm:px-6 lg:px-8 响应式容器\n"
    "• 颜色系统：primary-50/100/200/300/400/500/600/700/800/900十级色阶\n"
    "• 暗色模式兼容：dark:bg-surface-800 dark:text-surface-100",
    indent=True
)

add_para("3) JS业务逻辑", bold=True)
add_para(
    "Dashboard核心逻辑：\n"
    "• 使用TanStack Query 5进行服务端状态管理：useQuery获取文档列表，staleTime=5min自动缓存\n"
    "• 使用React useState管理本地UI状态（当前Tab、上传弹窗显示状态）\n"
    "• useUploadMutation：调用POST /api/v1/documents/upload，成功后invalidateQueries刷新列表\n"
    "• 统计卡片数据来自GET /api/v1/users/me/usage，每2分钟自动刷新\n"
    "• 文档卡片点击导航至/editor/:id路由\n"
    "• 错误处理：ErrorBoundary包裹，网络错误显示EmptyState组件",
    indent=True
)

# ── 2.2.2 DraggableAiAssistant ──
add_heading_styled("2.2.2  组件：DraggableAiAssistant组件", level=3)
add_para("1) 组件结构与DOM", bold=True)
add_para(
    "DraggableAiAssistant是一个全局可拖拽的AI对话浮窗组件（296行TypeScript代码），"
    "由浮动入口按钮（右下角紫色渐变圆形图标+在线状态指示灯）和展开的对话面板两部分组成。",
    indent=True
)
add_para(
    "展开面板DOM：\n"
    "div.fixed.z-50.w-[380px].bg-white.rounded-2xl.shadow-2xl.border（固定定位，可拖拽）\n"
    "├── Header（渐变背景，可拖拽手柄cursor-grab/grabbing）\n"
    "│   ├── DocMind助手 logo + 文档标题标签\n"
    "│   ├── 清空对话按钮\n"
    "│   └── 最小化按钮\n"
    "├── MessagesArea（flex-1 overflow-y-auto，消息列表）\n"
    "│   ├── Assistant消息（左对齐，白色气泡，圆角）\n"
    "│   └── User消息（右对齐，primary-500气泡）\n"
    "├── QuickActions（首条消息时显示：文档处理/文档摘要/格式转换/协作功能）\n"
    "└── InputArea（输入框 + 发送按钮）",
    indent=True, size=10
)

add_para("2) 样式与交互效果", bold=True)
add_para(
    "• 浮动按钮：w-14 h-14 bg-gradient-to-br from-primary-500 to-primary-700 rounded-full shadow-lg hover:scale-105\n"
    "• 对话面板：box-shadow: 0 20px 60px rgba(0,0,0,0.15)，520px高度，圆角2xl\n"
    "• 消息动画：loading状态使用3个bounce圆点（animation-delay: 0ms/150ms/300ms）\n"
    "• 过渡效果：所有hover/active状态使用transition-colors/transition-shadow\n"
    "• 可拖拽：mousedown/mousemove/mouseup事件 + clampPosition边界限制",
    indent=True
)

add_para("3) 组件核心逻辑", bold=True)
add_para(
    "DraggableAiAssistant使用React Hooks实现：\n"
    "• useState管理：isOpen（展开/收起）、messages（对话历史）、input（输入文字）、isLoading（加载状态）、position（面板位置）、isDragging（拖拽状态）\n"
    "• useRef管理：dragRef（拖拽起始坐标）、panelRef（面板DOM引用）、messagesEndRef（自动滚动锚点）、inputRef（输入框焦点）\n"
    "• useEffect：(1) 初始化位置（右下角）；(2) 新消息自动滚动到底部；(3) 展开时自动聚焦输入框\n"
    "• useCallback：(1) clampPosition限制面板不超出视口；(2) sendMessage发送消息并接收AI回复；(3) handleKeyDown支持Enter快捷发送；(4) clearChat清空对话\n"
    "• 消息流：用户输入→POST /api/v1/ai/chat→DeepSeek回复→更新messages状态",
    indent=True
)

# ── 2.2.3 ApiService ──
add_heading_styled("2.2.3  服务：ApiService服务层", level=3)
add_para("1) 服务架构", bold=True)
add_para(
    "前端API层采用Axios实例 + TanStack Query 5的组合架构：",
    indent=True
)
add_para(
    "// api.ts 核心结构\n"
    "const apiClient = axios.create({\n"
    "  baseURL: '/api/v1',\n"
    "  headers: { 'Content-Type': 'application/json' },\n"
    "});\n\n"
    "// 请求拦截器：自动注入Bearer Token\n"
    "apiClient.interceptors.request.use((config) => {\n"
    "  const token = getAccessToken();\n"
    "  if (token) config.headers.Authorization = `Bearer ${token}`;\n"
    "  return config;\n"
    "});\n\n"
    "// 响应拦截器：401自动刷新Token\n"
    "apiClient.interceptors.response.use(\n"
    "  (res) => res,\n"
    "  async (err) => {\n"
    "    if (err.response?.status === 401) {\n"
    "      const newToken = await refreshAccessToken();\n"
    "      err.config.headers.Authorization = `Bearer ${newToken}`;\n"
    "      return apiClient(err.config);\n"
    "    }\n"
    "    return Promise.reject(err);\n"
    "  }\n"
    ");",
    indent=True, size=10
)
add_para(
    "2) API模块化组织：\n"
    "• authApi：register/login/logout/refresh/oauthAccounts\n"
    "• documentApi：upload/list/get/update/delete/getContent/updateContent\n"
    "• aiApi：chat/proofread/summarize/rewrite/extract/convert/ocr/qa\n"
    "• searchApi：search/searchQA/buildIndex/deleteIndex\n"
    "• collaborationApi：createSession/getSession/invite/leave/remove\n"
    "• githubApi：getRepos/getContents/getReadme/getFile/importFile\n"
    "• userApi：getProfile/updateProfile/getUsage/getRecommendations/getAchievements\n"
    "• exportApi：export/compress/watermark/encrypt/merge",
    indent=True
)

# ── 2.2.4 i18n ──
add_heading_styled("2.2.4  管道：i18n国际化管道", level=3)
add_para("1) 国际化架构", bold=True)
add_para(
    "自定义i18n引擎（不使用第三方库），支持zh-CN和en-US两种语言，"
    "通过React Context在全局注入当前语言和切换函数。",
    indent=True
)
add_para(
    "// i18n.ts核心实现\n"
    "const I18nContext = createContext<I18nContextType>({});\n\n"
    "export function I18nProvider({ children }) {\n"
    "  const [locale, setLocale] = useState<'zh' | 'en'>(\n"
    "    () => localStorage.getItem('locale') || 'zh'\n"
    "  );\n"
    "  const t = useCallback((key: string) => {\n"
    "    return translations[locale][key] || key;\n"
    "  }, [locale]);\n"
    "  const toggleLocale = useCallback(() => {\n"
    "    setLocale(prev => prev === 'zh' ? 'en' : 'zh');\n"
    "  }, []);\n"
    "  return (\n"
    "    <I18nContext.Provider value={{ locale, t, toggleLocale }}>\n"
    "      {children}\n"
    "    </I18nContext.Provider>\n"
    "  );\n"
    "}",
    indent=True, size=10
)
add_para(
    "2) 词条管理：\n"
    "• zh.json：约500条中文词条，涵盖导航/页面/按钮/提示/错误信息\n"
    "• en.json：对应英文翻译，保持与zh.json相同的key结构\n"
    "• 使用方式：组件内调用 const { t } = useI18n()，通过 t('nav.dashboard') 获取翻译\n"
    "• 语言切换：Navbar右上角Switch to English/切换到中文按钮",
    indent=True
)

doc.add_page_break()

# ═══════════════════════════════════════════
# SECTION 3: 微服务设计与开发
# ═══════════════════════════════════════════

add_heading_styled("3  微服务设计与开发", level=1)

add_heading_styled("3.1  UML时序图设计", level=2)
add_para("以下PlantUML时序图展示核心业务场景的交互流程：", indent=True)

add_para("（1）用户认证流程（JWT双令牌机制）", bold=True)
add_para(
    "@startuml\n"
    "actor 用户\n"
    "participant \"React前端\" as FE\n"
    "participant \"FastAPI后端\" as BE\n"
    "database \"PostgreSQL\" as DB\n"
    "database \"Redis\" as Cache\n\n"
    "== 注册 ==\n"
    "用户 -> FE: 填写邮箱+密码\n"
    "FE -> BE: POST /auth/register\n"
    "BE -> DB: 检查邮箱唯一性\n"
    "BE -> BE: Argon2id哈希密码\n"
    "BE -> DB: INSERT users\n"
    "BE --> FE: 201 {id, email}\n\n"
    "== 登录 ==\n"
    "用户 -> FE: 输入凭证\n"
    "FE -> BE: POST /auth/login\n"
    "BE -> DB: SELECT user by email\n"
    "BE -> BE: Argon2id校验密码\n"
    "BE -> BE: 签发AccessToken(15min)+RefreshToken(7d)\n"
    "BE -> DB: INSERT refresh_tokens\n"
    "BE --> FE: 200 {access_token, refresh_token}\n"
    "FE -> FE: localStorage存储tokens\n\n"
    "== 请求认证 ==\n"
    "FE -> BE: GET /users/me (Authorization: Bearer xxx)\n"
    "BE -> BE: JWT中间件解码+校验\n"
    "BE --> FE: 200 {user data}\n\n"
    "== Token刷新 ==\n"
    "FE -> BE: POST /auth/refresh\n"
    "BE -> DB: 校验refresh_token+黑名单检查\n"
    "BE -> BE: 签发新AccessToken\n"
    "BE --> FE: 200 {new_access_token}\n"
    "@enduml",
    indent=True, size=9
)

add_para("（2）语义搜索→RAG问答流程", bold=True)
add_para(
    "@startuml\n"
    "actor 用户\n"
    "participant \"React前端\" as FE\n"
    "participant \"FastAPI\" as BE\n"
    "participant \"BCE Embedding\" as EMB\n"
    "participant \"FAISS Index\" as VEC\n"
    "participant \"DeepSeek LLM\" as LLM\n\n"
    "== 文档上传时：自动索引 ==\n"
    "FE -> BE: POST /documents/upload\n"
    "BE -> BE: 解析文档内容\n"
    "BE -> BE: 滑动窗口分块(chunk=1000, overlap=100)\n"
    "BE -> EMB: 文本向量化(768维)\n"
    "EMB --> BE: embeddings[]\n"
    "BE -> VEC: faiss.write_index(embeddings)\n\n"
    "== 混合搜索 ==\n"
    "用户 -> FE: 输入查询关键词\n"
    "FE -> BE: POST /search {mode: hybrid}\n"
    "BE -> EMB: 查询向量化\n"
    "BE -> VEC: faiss.search(top_k*2个候选)\n"
    "BE -> BE: BM25关键词评分\n"
    "BE -> BE: 加权融合: score = 0.7*vec + 0.3*bm25\n"
    "BE -> BE: 关键词高亮片段提取\n"
    "BE --> FE: 返回排序结果+双分数+高亮\n\n"
    "== RAG问答 ==\n"
    "用户 -> FE: 输入问题\n"
    "FE -> BE: POST /search/qa {question}\n"
    "BE -> VEC: 检索相关片段(top_k)\n"
    "BE -> LLM: 构建Prompt(系统+上下文+问题)\n"
    "LLM --> BE: 生成答案\n"
    "BE --> FE: 答案+引用来源\n"
    "@enduml",
    indent=True, size=9
)

add_heading_styled("3.2  关系数据库模式设计", level=2)
add_para("DocMind数据库包含16张核心表，以下为E-R关系图描述及关键表结构：", indent=True)

add_para("E-R关系概要：", bold=True)
add_para(
    "users (16 fields) ──< documents (20+ fields) ──< document_versions\n"
    "documents ──< collaboration_sessions ──< collaborators\n"
    "documents ──< ai_processing_jobs\n"
    "documents ──< tag_associations >── tags\n"
    "users ──< operation_logs\n"
    "users ──< oauth_accounts\n"
    "users ──< refresh_tokens\n"
    "users ── user_quotas (1:1)\n"
    "users ──< github_connections\n"
    "独立表: tier_definitions, enterprise_api_keys, cache_invalidation_keys",
    indent=True, size=10
)

add_para("核心表——users表SQL DDL：", bold=True)
add_para(
    "CREATE TABLE users (\n"
    "    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n"
    "    email VARCHAR(255) UNIQUE NOT NULL,\n"
    "    password_hash VARCHAR(255) NOT NULL,\n"
    "    display_name VARCHAR(100),\n"
    "    avatar_url VARCHAR(500),\n"
    "    tier VARCHAR(50) DEFAULT 'novice',\n"
    "    tier_expires_at TIMESTAMPTZ,\n"
    "    is_active BOOLEAN DEFAULT TRUE,\n"
    "    is_verified BOOLEAN DEFAULT FALSE,\n"
    "    is_superuser BOOLEAN DEFAULT FALSE,\n"
    "    preferences JSONB DEFAULT '{}',\n"
    "    quota_used_docs INTEGER DEFAULT 0,\n"
    "    quota_used_ai_calls INTEGER DEFAULT 0,\n"
    "    quota_used_storage_bytes BIGINT DEFAULT 0,\n"
    "    quota_period_start TIMESTAMPTZ,\n"
    "    last_login_at TIMESTAMPTZ,\n"
    "    created_at TIMESTAMPTZ DEFAULT NOW(),\n"
    "    updated_at TIMESTAMPTZ DEFAULT NOW()\n"
    ");",
    indent=True, size=10
)

add_para("核心表——documents表SQL DDL：", bold=True)
add_para(
    "CREATE TABLE documents (\n"
    "    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n"
    "    owner_id UUID NOT NULL REFERENCES users(id),\n"
    "    title VARCHAR(500) NOT NULL,\n"
    "    input_format VARCHAR(20) NOT NULL,\n"
    "    output_format VARCHAR(20),\n"
    "    mime_type VARCHAR(100),\n"
    "    file_size_bytes BIGINT,\n"
    "    page_count INTEGER,\n"
    "    char_count INTEGER,\n"
    "    status VARCHAR(20) DEFAULT 'processing',\n"
    "    storage_path VARCHAR(500),\n"
    "    parsed_content_path VARCHAR(500),\n"
    "    metadata JSONB DEFAULT '{}',\n"
    "    tags TEXT[],\n"
    "    folder VARCHAR(255),\n"
    "    is_deleted BOOLEAN DEFAULT FALSE,\n"
    "    deleted_at TIMESTAMPTZ,\n"
    "    current_version_id UUID,\n"
    "    checksum_sha256 VARCHAR(64),\n"
    "    created_at TIMESTAMPTZ DEFAULT NOW(),\n"
    "    updated_at TIMESTAMPTZ DEFAULT NOW()\n"
    ");",
    indent=True, size=10
)

add_para("其余14张表结构摘要：", bold=True)
other_tables = [
    ["document_versions", "version_number, content, change_summary, storage_path, created_at", "文档版本快照"],
    ["operation_logs", "user_id, document_id, action, action_category, details(JSONB)", "操作审计日志"],
    ["collaboration_sessions", "document_id, status(active/archived), created_by, created_at", "协作会话"],
    ["collaborators", "session_id, user_id, role(owner/editor/viewer), joined_at", "协作者管理"],
    ["ai_processing_jobs", "document_id, task_type, status, chunks_total, tokens_used, result", "AI异步任务队列"],
    ["oauth_accounts", "user_id, provider(github), provider_user_id, access_token(encrypted)", "OAuth绑定"],
    ["github_connections", "user_id, github_username, access_token(encrypted), avatar_url", "GitHub集成"],
    ["refresh_tokens", "user_id, token_hash, expires_at, is_revoked", "JWT刷新令牌"],
    ["tier_definitions", "name, max_documents, max_ai_calls, max_storage, price_monthly", "套餐定义"],
    ["user_quotas", "user_id, docs_used, ai_calls_used, storage_used, period_start", "用户配额"],
    ["tags", "name, color, owner_id", "文档标签"],
    ["tag_associations", "document_id, tag_id", "文档-标签关联"],
    ["enterprise_api_keys", "user_id, api_key_hash, name, rate_limit, is_active", "企业API密钥"],
    ["cache_invalidation_keys", "prefix, version, updated_at", "缓存失效控制"],
]
add_table(["表名", "关键字段", "用途"], other_tables)

add_heading_styled("3.3  范式数据的实体导入", level=2)
add_para(
    "DocMind采用3NF范式设计，所有表均满足第三范式要求：\n"
    "• 每个字段不可再分（1NF）：所有字段为原子类型\n"
    "• 非主键字段完全函数依赖于主键（2NF）：无部分依赖\n"
    "• 非主键字段不传递依赖于主键（3NF）：无传递依赖\n\n"
    "通过ORM层（SQLAlchemy 2.0）定义模型，使用Alembic进行数据库迁移管理。"
    "初始化数据（tier_definitions四档套餐）在应用启动时通过_seed_tiers()函数自动导入。"
    "测试环境使用SQLite + Dev Fallback模式，无需外部数据库依赖。",
    indent=True
)

add_para("种子数据——tier_definitions初始化INSERT：", bold=True)
add_para(
    "INSERT INTO tier_definitions VALUES\n"
    "('novice', 'Novice', 10, 20, 100000000, 5000000, 50000, '[\"pdf\",\"docx\",\"txt\",\"md\"]', false, false, false, 0, 0);\n"
    "('white_collar', 'White-Collar', 50, 100, 500000000, 25000000, 200000, '[\"pdf\",\"docx\",\"txt\",\"md\",\"png\",\"jpg\"]', true, true, false, 3, 9.99);\n"
    "('professional', 'Professional', 200, 500, 2000000000, 100000000, NULL, '[\"pdf\",\"docx\",\"txt\",\"md\",\"png\",\"jpg\",\"pptx\",\"xlsx\"]', true, true, false, 10, 29.99);\n"
    "('enterprise', 'Enterprise', 2147483647, 2147483647, 50000000000, 500000000, NULL, '[...]', true, true, true, NULL, 99.99);",
    indent=True, size=10
)

doc.add_page_break()

# ═══════════════════════════════════════════
# SECTION 4: 难点功能实现
# ═══════════════════════════════════════════

add_heading_styled("4  难点功能实现", level=1)

add_para("以下列举本项目技术难度较高的三个功能及其解决方案：", indent=True)

add_para("难点一：BCE Embedding模型的ONNX加载与跨平台部署", bold=True)
add_para(
    "问题描述：语义搜索需要加载530MB的BCE-Embedding-Base-v1嵌入模型。"
    "HuggingFace在Windows环境下存在SSL证书验证和符号链接支持问题，"
    "且PyTorch格式模型无法直接在无GPU环境高效运行。",
    indent=True
)
add_para(
    "解决方案：\n"
    "1. 通过ModelScope国内镜像下载ONNX格式模型（规避HF网络问题）\n"
    "2. 实现_OnnxEmbeddingModel自定义推理包装器，使用onnxruntime CPU执行\n"
    "3. 集成SentencePiece分词器处理XLMRoberta的tokenization\n"
    "4. 实现mean pooling + L2归一化的后处理管线\n"
    "5. 自动降级：优先尝试sentence_transformers（HF），失败后回退到ModelScope ONNX\n"
    "6. 创建modules.json和1_Pooling/config.json以兼容SentenceTransformer接口",
    indent=True
)

add_para("难点二：混合搜索的BM25+向量加权融合算法", bold=True)
add_para(
    "问题描述：纯向量搜索对专业术语和精确关键词匹配效果不佳；"
    "纯关键词搜索无法理解语义相似性。需要设计合理的分数融合策略。",
    indent=True
)
add_para(
    "解决方案：\n"
    "1. 实现CJK字符级和英文词级双模式BM25分词器（_tokenize函数）\n"
    "2. 简化BM25评分：使用TF饱和度公式 score = t / (t + k1 * (1 - b + b * dl / avgdl))\n"
    "3. BM25分数归一化：kw_norm = min(kw_score * 3.0, 1.0)，缩放到[0,1]区间\n"
    "4. 融合公式：fused_score = 0.7 * vector_score + 0.3 * kw_norm（可配置vector_weight）\n"
    "5. 先取top_k*2个向量候选，再融合排序，返回top_k个结果\n"
    "6. 搜索高亮：查询词匹配后提取15词窗口片段，用<< >>标记",
    indent=True
)

add_para("难点三：可拖拽AI助手的边界限制与性能优化", bold=True)
add_para(
    "问题描述：全局浮窗组件需要支持鼠标拖拽移动，同时保证不超出视口边界，"
    "并在最小化/展开状态间平滑过渡，且不能影响主页面性能。",
    indent=True
)
add_para(
    "解决方案：\n"
    "1. 实现clampPosition算法：实时计算面板宽高，限制x/y在[0, viewport-dim]区间\n"
    "2. 使用useRef存储拖拽起始位置，避免不必要的re-render\n"
    "3. useEffect + window事件监听mousemove/mouseup，确保拖拽流畅\n"
    "4. isDragging状态下切换cursor样式（grab↔grabbing）\n"
    "5. 最小化状态：渲染轻量浮动按钮（14x14圆形容器），无DOM挂载整个面板\n"
    "6. 初始化位置计算：useEffect中根据window.innerWidth/Height定位右下角\n"
    "7. 消息列表使用ref + scrollIntoView实现新消息自动滚动",
    indent=True
)

doc.add_page_break()

# ═══════════════════════════════════════════
# SECTION 5: 项目管理总结
# ═══════════════════════════════════════════

add_heading_styled("5  项目管理总结", level=1)

add_heading_styled("5.1  项目迭代周期分析", level=2)
add_para(
    "本项目采用敏捷迭代开发方式，共经历三个主要阶段：",
    indent=True
)
sprints = [
    ["Phase 1 (基础功能)", "2026.06.04—06.05", "6个commit",
     "多格式文档上传/解析、6种AI工具、实时协作、版本控制、GitHub集成、Dashboard UI重构、i18n国际化、Landing/Login/Register页面"],
    ["Phase 2 (功能增强)", "2026.06.05—06.09", "12个commit",
     "AI助手（可拖拽对话）、语义搜索（FAISS+BCE）、OCR识别（PaddleOCR+EasyOCR）、PDF导出增强（页码/目录/书签）、混合搜索（BM25+向量）、跨文档搜索、关键词高亮、语言检测、二维码识别"],
    ["Phase 3 (测试完善)", "2026.06.09", "3个commit",
     "单元测试19个、集成测试30个、API覆盖率测试18个、合计67项全通过（0 skipped, 0 warnings）"],
]
add_table(["阶段", "时间", "提交数", "主要成果"], sprints)

# 插入GitHub Pulse活跃度截图
add_para("图1：GitHub项目活跃度（Pulse）", bold=True, size=10)
img_path = os.path.join(SCREENSHOT_DIR, "screenshot_pulse.png")
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading_styled("5.2  项目成员贡献分析", level=2)
add_para(
    "本项目为单人全栈项目，独立完成前端React、后端FastAPI、AI服务集成及测试的全部开发工作。",
    indent=True
)
add_para(
    "贡献统计（通过Git记录）：\n"
    "• 总提交数：58+ commits\n"
    "• Python源文件：92个（含API路由、服务层、模型、Schema、AI引擎、Worker）\n"
    "• TypeScript源文件：53个（含页面、组件、工具库、类型定义、国际化词条）\n"
    "• 代码行数：约8000+行（Python ~5000行 + TypeScript ~3000行）\n"
    "• 测试覆盖：67项（19单元 + 30集成 + 18 API）",
    indent=True
)

# 插入贡献者截图
add_para("图2：GitHub贡献者统计（Contributors）", bold=True, size=10)
img_path = os.path.join(SCREENSHOT_DIR, "screenshot_contributors.png")
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading_styled("5.3  项目代码量分析", level=2)
code_stats = [
    ["Python", "src/api/v1/ (12 routes)", "~1500行", "RESTful API端点"],
    ["Python", "src/services/ (10 modules)", "~1800行", "业务逻辑层"],
    ["Python", "src/ai/ (5 modules)", "~900行", "AI引擎"],
    ["Python", "src/models/ + schemas/", "~500行", "ORM模型+Pydantic"],
    ["TypeScript", "frontend/src/pages/ (12 pages)", "~1200行", "页面组件"],
    ["TypeScript", "frontend/src/components/ (26 comps)", "~1200行", "可复用组件"],
    ["TypeScript", "frontend/src/lib/ + types/", "~400行", "工具库+类型"],
    ["Python/TS", "tests/ (6 files)", "~400行", "测试代码"],
]
add_table(["语言", "模块", "代码量", "说明"], code_stats)

# 插入仓库和提交截图
add_para("图3：GitHub仓库主页", bold=True, size=10)
img_path = os.path.join(SCREENSHOT_DIR, "screenshot_repo_main.png")
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_para("图4：GitHub提交历史（58+ commits）", bold=True, size=10)
img_path = os.path.join(SCREENSHOT_DIR, "screenshot_commits.png")
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ═══════════════════════════════════════════
# SECTION 6: 项目产品展示
# ═══════════════════════════════════════════

add_heading_styled("6  项目产品展示", level=1)

add_para("以下为DocMind平台主要页面的功能说明（实际截图请另行插入）：", indent=True)

screenshots = [
    ["Landing着陆页", "展示DocMind品牌定位、核心功能卡片、四档定价方案、CTA按钮"],
    ["Dashboard工作台", "统计卡片（文档数/AI调用/存储）、文档列表（按Tab筛选）、上传按钮、AI助手浮动入口"],
    ["AI助手对话", "可拖拽对话面板、快速提问标签、DeepSeek多轮对话、文档上下文感知"],
    ["AI工具面板", "校对/总结/改写/翻译/转换等9种AI工具、参数配置、实时进度"],
    ["语义搜索面板", "双Tab（搜索/问答）、混合排序、关键词高亮、RAG答案+来源标注"],
    ["PPTX幻灯片查看", "16页幻灯片逐页渲染、文本提取显示"],
    ["版本历史", "版本时间线、回滚操作、Diff内容对比"],
    ["协作管理", "创建协作会话、邀请协作者、设置角色权限"],
]
add_table(["页面/功能", "说明"], screenshots)

add_para(
    "注：详细的产品演示请参见DocMind_项目汇报.pptx（16页完整汇报幻灯片）和30秒以上的项目展示视频。",
    indent=True
)

# ═══════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════

doc.save(OUTPUT)
print(f"Report saved to: {OUTPUT}")
print(f"File size: {os.path.getsize(OUTPUT) / 1024:.1f} KB")
